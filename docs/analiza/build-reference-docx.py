"""Build a properly styled reference.docx for Pandoc academic conversion.

Stiluri:
- Normal: Times New Roman 12pt, line spacing 1.15, justified
- Heading 1: Times 16pt bold, space before 24pt, after 12pt, page break before
- Heading 2: Times 14pt bold, space before 18pt, after 6pt
- Heading 3: Times 12pt bold, space before 12pt, after 6pt
- Margins: 2.5cm all sides
- Page numbers in footer (centered)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import sys

REF = "/Users/octav/MODBD/docs/analiza/reference.docx"

doc = Document(REF)

# ---- 1. Set page margins (2.5cm all sides) ----
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---- 2. Helper for setting style properties ----
def set_font(style, name="Times New Roman", size=12, bold=False, italic=False):
    """Force font on all relevant style.font.* layers (style XML + rFonts)."""
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    # Force East Asian + complex script font too (Word fallback)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)

def set_paragraph_format(style, before=0, after=0, line_spacing=None,
                         alignment=None, keep_with_next=False):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if alignment is not None:
        pf.alignment = alignment
    pf.keep_with_next = keep_with_next

# ---- 3. Normal style: Times 12pt, line-spacing 1.15, justified ----
normal = doc.styles["Normal"]
set_font(normal, "Times New Roman", 12)
set_paragraph_format(normal, before=0, after=8, line_spacing=1.15,
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ---- 4. Heading 1: 16pt bold, big spacing, keep with next ----
h1 = doc.styles["Heading 1"]
set_font(h1, "Times New Roman", 16, bold=True)
set_paragraph_format(h1, before=24, after=12, line_spacing=1.15,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=True)
# Color black (override default blue)
h1.font.color.rgb = RGBColor(0, 0, 0)

# ---- 5. Heading 2: 14pt bold ----
h2 = doc.styles["Heading 2"]
set_font(h2, "Times New Roman", 14, bold=True)
set_paragraph_format(h2, before=18, after=6, line_spacing=1.15,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=True)
h2.font.color.rgb = RGBColor(0, 0, 0)

# ---- 6. Heading 3: 12pt bold ----
h3 = doc.styles["Heading 3"]
set_font(h3, "Times New Roman", 12, bold=True)
set_paragraph_format(h3, before=12, after=6, line_spacing=1.15,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=True)
h3.font.color.rgb = RGBColor(0, 0, 0)

# ---- 7. Title style (used by frontmatter "title" field) ----
try:
    title = doc.styles["Title"]
    set_font(title, "Times New Roman", 20, bold=True)
    set_paragraph_format(title, before=0, after=24, line_spacing=1.15,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    title.font.color.rgb = RGBColor(0, 0, 0)
except KeyError:
    pass

# ---- 8. Subtitle (frontmatter "subtitle") ----
try:
    subtitle = doc.styles["Subtitle"]
    set_font(subtitle, "Times New Roman", 14, italic=True)
    set_paragraph_format(subtitle, before=0, after=12, line_spacing=1.15,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle.font.color.rgb = RGBColor(80, 80, 80)
except KeyError:
    pass

# ---- 9. Author + Date (custom paragraphs from frontmatter) ----
for sname in ["Author", "Date"]:
    try:
        s = doc.styles[sname]
        set_font(s, "Times New Roman", 12)
        set_paragraph_format(s, before=0, after=4, line_spacing=1.15,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except KeyError:
        pass

# ---- 10. Source code blocks (Pandoc "Source Code") ----
try:
    src = doc.styles["Source Code"]
    set_font(src, "Menlo", 10)
except KeyError:
    pass

# ---- 11. Table style: borders + Times font ----
try:
    tbl = doc.styles["Table Grid"]
    # Borders are already set in Table Grid; ensure font
    set_font(tbl, "Times New Roman", 11)
except KeyError:
    pass

# ---- 12. Page numbers in footer (centered) ----
def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Clear existing runs
    for r in list(p.runs):
        r.text = ""
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    # Ensure run font is Times 11pt
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

for section in doc.sections:
    add_page_number(section)

# ---- 13. Save ----
doc.save(REF)
print(f"Saved styled reference.docx → {REF}")
print("Sizes applied: Normal 12, H1 16, H2 14, H3 12, line-spacing 1.15, margins 2.5cm")
