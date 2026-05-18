"""Generate Echipa.docx with proper academic formatting.

Stiluri:
- Title: Times 16pt bold, centered
- Heading 1: Times 14pt bold (sectiuni majore)
- Heading 2: Times 13pt bold (persoane)
- Heading 3: Times 12pt bold (sub-categorii task)
- Normal: Times 12pt justified
- Bullet list: Times 12pt indented
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/octav/MODBD/docs/analiza/output/NUME_ECHIPA_Oprinoiu_Octavian_Echipa.docx"

doc = Document()

# Page margins 2.5cm
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Force Times New Roman everywhere
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16, bold=True)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_font(r, size=12, italic=True, color=RGBColor(80, 80, 80))
    return p

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=14, bold=True)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=13, bold=True)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=12, bold=True, italic=True)
    return p

def add_para(text, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75)
    for r in p.runs:
        set_font(r)
    if not p.runs:
        r = p.add_run(text)
        set_font(r)
    else:
        p.runs[0].text = text
        for r in p.runs:
            set_font(r)
    return p

def add_responsabilitate(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Responsabilitate principala: ")
    set_font(r, bold=True)
    r2 = p.add_run(text)
    set_font(r2, italic=True)
    return p


# === DOCUMENT CONTENT ===

add_title("MODBD - Componenta echipei si distributia task-urilor")
add_subtitle("Proiect: Metode de Optimizare si Distribuire in Baze de Date (MODBD) "
             "- FMI Universitatea Bucuresti, anul universitar 2025-2026")
add_subtitle("Echipa: <<NUME_ECHIPA>>")


add_h1("Componenta echipei")
add_bullet("Stefan Magureanu")
add_bullet("Octavian Oprinoiu")
add_bullet("Andrei Pitoiu")


add_h1("Distributia task-urilor")


# ===== OCTAV =====
add_h2("Octavian Oprinoiu")
add_responsabilitate("Pregatirea datelor + Backend baze de date + Build automation.")

add_h3("Pregatirea datelor sursa")
for b in [
    "Analiza bazei de date OLTP sursa (SQL Server, aproximativ 280 tabele, 52 milioane randuri); identificarea subset-ului coerent pentru proiect.",
    "Selectia subset-ului final: 15 tabele (12 entitati independente + 3 relatii many-to-many), 95 coloane pastrate din 202 initiale, 10 clienti reprezentativi acoperind 5 zone, 3 valute si 6 ani de istoric.",
    "Scripturile de extragere a datelor din baza SQL Server sursa.",
    "Anonimizarea datelor sensibile (cod_client, denumire_client, cod_agent, email-uri) cu mapping determinist, asistata de un sistem AI pentru aplicarea consistenta in toate cele 15 fisiere CSV.",
    "Exportul celor 15 fisiere CSV finale, validate pentru integritate referentiala.",
]:
    add_bullet(b)

add_h3("Implementarea bazei de date distribuite (Modul 2)")
for b in [
    "Crearea celor 3 PDB-uri Oracle (DISTRIBUTIE, CATALOG, VANZARI) plus tablespaces.",
    "Crearea utilizatorilor aplicativi (sgbd_distributie, sgbd_catalog, sgbd_vanzari) plus role sgbd_role cu grant-urile complete.",
    "DDL pentru toate cele 18 tabele master si fragmente: 8 tabele in DISTRIBUTIE (CRM / comercial); 4 lookups plus ITEMS_CORE si ITEMS_EXTRA in CATALOG (fragmentare verticala); 4 fragmente fizice orizontale in VANZARI (FISE_CLIENTI_RO / EXT plus LINII_DOC_RO / EXT).",
    "Implementarea fragmentarii orizontale prin split pe Moneda la momentul incarcarii datelor.",
    "Incarcarea datelor din CSV-uri prin external tables (ORACLE_LOADER).",
    "Layerul de transparenta: view-uri UNION ALL (V_FISE_CLIENTI, V_LINII_DOC), V_ITEMS (JOIN intre CORE si EXTRA), plus 8 triggere INSTEAD OF pentru rutarea DML.",
    "Sincronizarea relatiilor replicate: 7 MV LOGs pe tabelele master plus 7 Materialized Views REFRESH FAST in VANZARI plus DBMS_SCHEDULER job la 60 de secunde.",
    "DB Links cross-PDB (lnk_distributie, lnk_catalog) si 4 chei externe cross-PDB catre MV-urile replicate.",
    "Trigger AFTER STATEMENT pentru coerenta sumelor document vs sumelor de pe linii (constrangere semantica cu agregat, conform cap. 3.4 clasa 3 din curs).",
    "8 indecsi pe fact tables plus DBMS_STATS.GATHER_SCHEMA_STATS pentru optimizatorul bazat pe cost.",
]:
    add_bullet(b)

add_h3("Build automation")
for b in [
    "Script master pentru replicarea automata a intregii baze de date distribuite intr-un container Docker fresh.",
    "Script de consolidare a celor 4 livrabile SQL (Setup.sql, Demo_Schema.sql, Demo_Queries.sql, Sursa.txt) pornind de la cele 18 scripturi sursa modulare.",
]:
    add_bullet(b)


# ===== ANDREI =====
add_h2("Andrei Pitoiu")
add_responsabilitate("Raport de Analiza + Cerere SQL si Optimizare + Validare end-to-end.")

add_h3("Raportul de Analiza (Modul 1)")
for b in [
    "Redactarea integrala a raportului de analiza (NUME_ECHIPA_Oprinoiu_Octavian_Analiza.docx).",
    "Descrierea modelului de afaceri (distributie B2B fashion in Romania, Cehia si Slovacia) si obiectivele aplicatiei distribuite.",
    "Diagrama Entitate-Relatie a bazei OLTP initiale (notatie Chen, 12 entitati independente plus 3 relatii M:N, normalizare FN3).",
    "Diagrama conceptuala globala (Crow's foot, atribute complete cu marcare PK / UK / FK si cardinalitati).",
    "Descrierea modului de distribuire (3 PDB-uri intr-un singur CDB Oracle).",
    "Argumentarea fragmentarii orizontale primare pe FISE_CLIENTI (predicat pe Moneda) plus aplicarea algoritmului COM_MIN.",
    "Argumentarea fragmentarii orizontale derivate pe LINII_DOC prin semijoin cu owner-ul FISE_CLIENTI.",
    "Argumentarea fragmentarii verticale pe MS_ITEMS prin algoritmul BEA (matricea VA, calculul afinitatilor, algoritmul PART, calculul z).",
    "Verificarea corectitudinii celor 3 fragmentari (completitudine, reconstructie, disjunctie).",
    "Argumentarea deciziei de replicare (criterii: volum, raport citire / scriere, locatia join-urilor).",
    "Schemele conceptuale locale pentru cele 3 PDB-uri.",
    "Lista exhaustiva a constrangerilor de integritate (unicitate locala / globala, chei primare, chei externe locale plus cross-PDB, validari domain si agregate).",
]:
    add_bullet(b)

add_h3("Cerere SQL complexa + optimizare (Modul 2, 1.5p)")
for b in [
    'Formularea in limbaj natural a cererii complexe ("Top 10 agenti dupa valoare vanduta in 2024, defalcat pe zona comerciala si categorie de produs").',
    "Implementarea SQL a cererii (touch in toate cele 3 PDB-uri prin DB link plus view-uri de transparenta plus MV-uri replicate, 8 join-uri si doua agregari).",
    "Cele 3 planuri de executie EXPLAIN PLAN comparate: RBO (hint /*+ RULE */) - abordare euristica; CBO default (cu statistici) - echivalent System R; CBO cu /*+ DRIVING_SITE */ - echivalent SDD-1.",
    "Analiza comparativa a cost-urilor, descrierea diferentelor de strategie, motivarea celei mai eficiente variante.",
]:
    add_bullet(b)

add_h3("Validare end-to-end")
for b in [
    "Suite de teste end-to-end care confirma functionarea completa a bazei distribuite: counts globale, V_FISE = suma fragmentelor, INSTEAD OF trigger ruteaza corect pe predicatul de fragmentare, FK cross-PDB respinge clientii inexistenti, sincronizarea MV propaga delta dupa refresh manual.",
]:
    add_bullet(b)


# ===== STEFAN =====
add_h2("Stefan Magureanu")
add_responsabilitate("Aplicatie Front-end (Modul 3).")

add_h3("Implementare aplicatie (Modul 3)")
for b in [
    "Designul arhitecturii aplicatiei: alegerea stack-ului tehnologic, pattern-ul conexiunilor catre cele 3 PDB-uri Oracle.",
    "API / backend pentru aplicatie - layerul de comunicare cu cele 3 PDB-uri Oracle (driver, pool de conexiuni, mapping query-uri).",
    "Modul CRUD pentru BD-urile locale - introducerea si gestiunea informatiilor in cele 3 PDB-uri (cerinta 1 din Modul 3, 3 puncte).",
    "Modul de vizualizare la nivelul BD globale - aplicatia se conecteaza la PDB-ul VANZARI si opereaza prin view-urile de transparenta ca si cum datele nu ar fi distribuite la nivelul retelei (cerinta 2, 1 punct).",
    "Vizualizarea la nivel global a operatiilor LMD locale - reflectarea in interfata a modificarilor facute direct in fragmentele orizontale (RO / EXT), fragmentele verticale (CORE / EXTRA) si relatiile master, prin view-urile de transparenta si refresh-ul MV (cerinta 3, 2 puncte).",
    "Verificarea propagarii LMD global catre local - operatiile facute prin view-urile UNION ALL (V_FISE_CLIENTI, V_LINII_DOC) si V_ITEMS sunt rutate corect catre fragmentele fizice prin trigger-ele INSTEAD OF (cerinta 4, 3 puncte).",
    "Realizarea capturilor de ecran din aplicatie pentru livrabilul NUME_ECHIPA_Oprinoiu_Octavian_Aplicatie.docx.",
]:
    add_bullet(b)


# ===== COMUNE =====
add_h1("Task-uri comune (impartite intre toti membrii echipei)")
for b in [
    "Cele 6 diagrame Mermaid din raport (E-R OLTP, conceptuala globala, topologia de distributie, 3 scheme conceptuale locale per PDB) - design, iterare, polish.",
    "Coordonarea proiectului - sincronizari de progres, decizii arhitecturale, prioritizare task-uri.",
    "Pregatirea prezentarii orale - repetitie demo, anticipare intrebari din curs si laborator, polish final inainte de submisie.",
]:
    add_bullet(b)


# Save
doc.save(OUT)
print(f"Generat: {OUT}")

import os
size_kb = os.path.getsize(OUT) // 1024
print(f"Marime: {size_kb} KB")
